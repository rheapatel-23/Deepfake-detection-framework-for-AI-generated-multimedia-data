import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
MD_FILE = r"c:\Users\rheap\Desktop\DeepFake Multimedia Detection\PROJECT_REPORT.md"
DOCX_FILE = r"c:\Users\rheap\Desktop\DeepFake Multimedia Detection\Deepfake_Detection_Project_Report.docx"
IMAGE_DIR = r"C:\Users\rheap\.gemini\antigravity\brain\785204f9-d653-47ae-8be5-9ca3540fcae9"

def add_markdown_paragraph(doc, text, style='BodyText'):
    """Parses simple markdown bold (**text**) and adds to docx."""
    p = doc.add_paragraph(style=style)
    # Split by **
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def convert_to_docx():
    doc = Document()
    
    # Title Style
    style = doc.styles['Title']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(24)
    font.color.rgb = RGBColor(0, 51, 102) # Dark Blue

    # Heading 1 Style
    style = doc.styles['Heading 1']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(18)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Read Markdown
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    table_mode = False
    table_header = []
    table_rows = []
    
    for line in lines:
        line = line.strip()
        
        # Images
        if "confusion_matrix.png" in line:
            img_path = os.path.join(IMAGE_DIR, "confusion_matrix.png")
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(4))
                last_p = doc.paragraphs[-1] 
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if "roc_curve.png" in line:
            img_path = os.path.join(IMAGE_DIR, "roc_curve.png")
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(4))
                last_p = doc.paragraphs[-1] 
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Headers
        if line.startswith('# '):
            doc.add_paragraph(line[2:], style='Title')
        elif line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading 1')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:], style='Heading 2')
            
        # Lists
        elif line.startswith('- '):
            add_markdown_paragraph(doc, line[2:], style='List Bullet')
            
        # Tables (Simple detection)
        elif line.startswith('|'):
            if '---' in line:
                continue
            cols = [c.strip() for c in line.split('|') if c.strip()]
            if not table_mode:
                table_mode = True
                table_header = cols
            else:
                table_rows.append(cols)
        
        # Code Blocks
        elif line.startswith('```'):
            continue # Skip markers
            
        # Normal Text
        elif line:
            if table_mode:
                 # Flush table if we hit text
                 table = doc.add_table(rows=1, cols=len(table_header))
                 table.style = 'Table Grid'
                 hdr_cells = table.rows[0].cells
                 for i, h in enumerate(table_header):
                     hdr_cells[i].text = h
                     hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                     
                 for row in table_rows:
                     row_cells = table.add_row().cells
                     for i, c in enumerate(row):
                         if i < len(row_cells):
                             row_cells[i].text = c
                 
                 doc.add_paragraph() # Spacer
                 table_mode = False
                 table_rows = []
                 table_header = []
                 add_markdown_paragraph(doc, line)
            else:
                 add_markdown_paragraph(doc, line)
        else:
            # Empty line
            if table_mode:
                 # Flush table
                 table = doc.add_table(rows=1, cols=len(table_header))
                 table.style = 'Table Grid'
                 hdr_cells = table.rows[0].cells
                 for i, h in enumerate(table_header):
                     hdr_cells[i].text = h
                     hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                     
                 for row in table_rows:
                     row_cells = table.add_row().cells
                     for i, c in enumerate(row):
                         if i < len(row_cells):
                             row_cells[i].text = c
                 doc.add_paragraph()
                 table_mode = False
                 table_rows = []
                 table_header = []

    print(f"Saved Word document to: {DOCX_FILE}")
    doc.save(DOCX_FILE)

if __name__ == "__main__":
    convert_to_docx()
